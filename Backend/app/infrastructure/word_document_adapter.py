from io import BytesIO
from docx import Document
from docx.shared import Inches

from app.domain.IDocumentGenerator import IDocumentGenerator
from app.domain.entities import Quiz

class WordDocumentAdapter(IDocumentGenerator):
    
    def create_exam_document(self, quiz_data: Quiz) -> BytesIO:
        document = Document()
        document.add_heading('Examen de Evaluación', level=1)
        document.add_paragraph()

        for i, item in enumerate(quiz_data.quiz, start=1):
            document.add_heading(f'Pregunta {i}: {item.question}', level=2)
            for idx, option in enumerate(item.answer_options):
                letter = chr(ord('A') + idx)
                p = document.add_paragraph()
                p.add_run(f'{letter}) {option.text}')
                p.paragraph_format.left_indent = Inches(0.5)
            document.add_paragraph()

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def create_answer_key_document(self, quiz_data: Quiz) -> BytesIO:
        document = Document()
        document.add_heading('Hoja de Respuestas', level=1)
        document.add_paragraph()

        for i, item in enumerate(quiz_data.quiz, start=1):
            for idx, option in enumerate(item.answer_options):
                if option.is_correct:
                    correct_letter = chr(ord('A') + idx)
                    p = document.add_paragraph()
                    p.add_run(f'{i}. {correct_letter} - ').bold = True
                    p.add_run(option.text)
                    p_rationale = document.add_paragraph()
                    p_rationale.add_run(f'Justificación: {option.rationale}').italic = True
                    p_rationale.paragraph_format.left_indent = Inches(0.5)
                    break
    
        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream